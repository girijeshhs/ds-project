from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "report_assets"
ASSETS.mkdir(exist_ok=True)


def generate_dataset():
    np.random.seed(42)
    first_half_usage = np.random.randint(180, 320, size=180)
    second_half_usage = np.random.randint(260, 420, size=180)
    telecom_usage_360 = np.concatenate((first_half_usage, second_half_usage))
    telecom_usage_12x30 = telecom_usage_360.reshape(12, 30)

    months = [
        "January",
        "February",
        "March",
        "April",
        "May",
        "June",
        "July",
        "August",
        "September",
        "October",
        "November",
        "December",
    ]
    day_columns = [f"Day_{day}" for day in range(1, 31)]
    telecom_df = pd.DataFrame(telecom_usage_12x30, index=months, columns=day_columns)
    telecom_df["Monthly_Total"] = telecom_df.sum(axis=1)
    return telecom_usage_360, telecom_usage_12x30, telecom_df


def save_line_chart(telecom_df):
    chart_path = ASSETS / "monthly_telecom_usage.png"
    plt.figure(figsize=(10, 5))
    plt.plot(
        telecom_df.index,
        telecom_df["Monthly_Total"],
        marker="o",
        linewidth=2.5,
        color="teal",
    )
    plt.title("Monthly Telecom Usage Total", fontsize=15, weight="bold")
    plt.xlabel("Month")
    plt.ylabel("Total Usage")
    plt.xticks(rotation=45)
    plt.grid(True, linestyle="--", alpha=0.6)
    plt.tight_layout()
    plt.savefig(chart_path, dpi=220, bbox_inches="tight")
    plt.close()
    return chart_path


def save_dataframe_preview(telecom_df):
    preview_path = ASSETS / "dataframe_preview.png"
    preview_df = telecom_df.iloc[:5, :7].copy()
    preview_df.columns = [
        "Day_1",
        "Day_2",
        "Day_3",
        "Day_4",
        "Day_5",
        "Day_6",
        "Monthly_Total",
    ]

    fig, ax = plt.subplots(figsize=(12, 3.8))
    ax.axis("off")
    ax.set_title("Telecom DataFrame Preview", fontsize=15, weight="bold", pad=12)
    table = ax.table(
        cellText=preview_df.values,
        rowLabels=preview_df.index,
        colLabels=preview_df.columns,
        cellLoc="center",
        loc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.5)
    plt.tight_layout()
    plt.savefig(preview_path, dpi=220, bbox_inches="tight")
    plt.close()
    return preview_path


def save_console_output(telecom_usage_360, telecom_usage_12x30, telecom_df):
    output_path = ASSETS / "console_output.png"
    lines = [
        "Program Output Summary",
        "",
        f"Telecom usage array shape: {telecom_usage_360.shape}",
        f"Reshaped telecom dataset shape: {telecom_usage_12x30.shape}",
        "",
        "Monthly totals:",
    ]
    for month, total in telecom_df["Monthly_Total"].items():
        lines.append(f"{month:<10} : {int(total)}")

    fig, ax = plt.subplots(figsize=(8.5, 6.5))
    ax.axis("off")
    ax.text(
        0.02,
        0.98,
        "\n".join(lines),
        va="top",
        ha="left",
        family="monospace",
        fontsize=11,
        bbox={"facecolor": "#f4f6f8", "edgecolor": "#c7d0d9", "boxstyle": "round,pad=0.8"},
    )
    plt.tight_layout()
    plt.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close()
    return output_path


def save_workflow_diagram():
    diagram_path = ASSETS / "workflow_diagram.png"
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")

    boxes = [
        (5, 9, "Create Two Arrays\n(180 + 180 values)"),
        (5, 7.2, "Use concatenate()\nMake 360-value dataset"),
        (5, 5.4, "Use reshape(12, 30)\nFix telecom data shape"),
        (5, 3.6, "Convert to Pandas DataFrame\nAdd month/day labels"),
        (5, 1.8, "Calculate Monthly Totals\nPlot Line Graph"),
    ]

    for x, y, label in boxes:
        ax.text(
            x,
            y,
            label,
            ha="center",
            va="center",
            fontsize=12,
            weight="bold",
            bbox={"boxstyle": "round,pad=0.6", "facecolor": "#e8f4f8", "edgecolor": "#2d6f73", "linewidth": 1.5},
        )

    for start_y, end_y in [(8.35, 7.85), (6.55, 6.05), (4.75, 4.25), (2.95, 2.45)]:
        ax.annotate(
            "",
            xy=(5, end_y),
            xytext=(5, start_y),
            arrowprops={"arrowstyle": "->", "lw": 2, "color": "#2d6f73"},
        )

    ax.set_title("Diagrammatic Representation of the Code Workflow", fontsize=16, weight="bold", pad=20)
    plt.tight_layout()
    plt.savefig(diagram_path, dpi=220, bbox_inches="tight")
    plt.close()
    return diagram_path


def set_page_border(section):
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")

    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "16")
        border.set(qn("w:color"), "7F8C8D")
        pg_borders.append(border)

    sect_pr.append(pg_borders)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_page_border(section)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_page_heading(doc, text):
    heading = doc.add_paragraph(style="Heading 1")
    run = heading.add_run(text)
    run.bold = True


def add_paragraph(doc, text, bold=False, center=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(10)
    para.paragraph_format.line_spacing = 1.5
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.5
    para.add_run(text)


def add_numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.5
    para.add_run(text)


def add_image(doc, image_path, width=6.0, caption=None):
    doc.add_picture(str(image_path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(caption)
        caption_run.italic = True


def add_code_block(doc, code_text):
    for line in code_text.splitlines():
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)


def page_break(doc):
    doc.add_page_break()
    section = doc.sections[-1]
    set_page_border(section)


def build_document(telecom_usage_360, telecom_usage_12x30, telecom_df, chart_path, preview_path, output_path, diagram_path):
    doc = Document()
    style_document(doc)

    code_text = (ROOT / "telecom_data_shape_fix.py").read_text()
    peak_month = telecom_df["Monthly_Total"].idxmax()
    peak_total = int(telecom_df["Monthly_Total"].max())
    min_month = telecom_df["Monthly_Total"].idxmin()
    min_total = int(telecom_df["Monthly_Total"].min())

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("TELECOM DATA SHAPE FIX USING NUMPY AND PANDAS")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = "Times New Roman"

    for text in [
        "",
        "Mini Project Report",
        "",
        "Student Name: ______________________________",
        "Register Number: ______________________________",
        "Subject Name: Data Science",
        "Faculty Name: ______________________________",
        "College Name: ______________________________",
    ]:
        add_paragraph(doc, text, center=True)

    page_break(doc)

    add_page_heading(doc, "Abstract")
    add_paragraph(doc, "This mini-project explains how NumPy and Pandas can be used to correct the shape of telecom usage data and make it suitable for simple analysis. In many practical cases, telecom values are first collected as a long one-dimensional array. While this raw format stores the values, it does not present them in a way that is easy to study month by month or day by day. This report demonstrates how such raw data can be reorganized into a more meaningful academic example.")
    add_paragraph(doc, "A telecom dataset of 360 values is created to represent daily usage records across a year. The values are first handled as separate arrays and then unified into one full dataset using NumPy concatenate(). After that, reshape() is applied to convert the one-dimensional data into a 12 x 30 structure. This arrangement is useful because it allows each row to represent a month and each column to represent a day, making the dataset much easier to read.")
    add_paragraph(doc, "Once the shape is corrected, the reshaped array is converted into a Pandas DataFrame with proper labels. Monthly totals are then calculated and displayed so that the usage level of each month can be compared clearly. Finally, a Matplotlib line graph is used to visualize the telecom consumption trend. The overall project shows that data shape correction is a basic but essential step in telecom analytics because meaningful analysis becomes possible only after the data is properly arranged.")

    page_break(doc)

    add_page_heading(doc, "Problem Statement")
    add_paragraph(doc, "Telecom systems produce a large amount of daily data from calls, messages, and internet usage. In many situations, this information is collected continuously and stored as raw numerical values without a proper table structure. When data is available only as a long array, it becomes difficult for students or analysts to understand which values belong to which month, how much usage happened during a particular period, or when the network experienced higher demand.")
    add_paragraph(doc, "A raw list of values may still contain useful information, but without the correct shape it is not easy to interpret. For example, a 360-length array can represent a full year of 30-day months, but that meaning is not obvious unless the values are arranged properly. If the data is not reshaped into rows and columns, direct comparison between months becomes confusing and trend analysis becomes less reliable.")
    add_paragraph(doc, "In this project, the telecom dataset begins as a one-dimensional array of 360 values. The main problem is to convert this raw array into a meaningful 12 x 30 format, combine separate usage arrays into one unified dataset, calculate monthly totals, and visualize the result clearly. Solving this problem helps convert unorganized data into a form that supports analysis, reporting, and better understanding of telecom network behavior.")

    page_break(doc)

    add_page_heading(doc, "Objectives")
    objectives = [
        "To create a telecom usage dataset containing 360 values.",
        "To reshape the raw dataset into a 12 x 30 monthly structure using NumPy.",
        "To use concatenate() to merge separate telecom usage arrays.",
        "To convert the reshaped array into a Pandas DataFrame.",
        "To calculate monthly totals for comparison and trend study.",
        "To generate a line graph using Matplotlib.",
        "To understand how data shape correction supports telecom data analysis.",
        "To identify periods of low and high network load using monthly totals.",
        "To present the complete result in a neat, report-ready format.",
    ]
    for item in objectives:
        add_bullet(doc, item)
    add_paragraph(doc, "These objectives are designed to cover the full workflow from raw data creation to interpretation. Instead of focusing only on coding syntax, the project also emphasizes presentation, trend observation, and real-world telecom understanding. This makes the work more suitable for a data science assignment where both technical steps and final reporting are important.")

    page_break(doc)

    add_page_heading(doc, "Tools and Technologies Used")
    tools = [
        "Python for program development and execution.",
        "NumPy for array creation, concatenation, and reshaping.",
        "Pandas for DataFrame conversion and monthly total calculation.",
        "Matplotlib for plotting the line graph.",
        "Microsoft Word for professional report presentation.",
    ]
    for item in tools:
        add_bullet(doc, item)
    add_paragraph(doc, "Each tool used in this project has a clear role. NumPy handles the core array operations, Pandas improves readability by introducing a table structure, and Matplotlib makes the result visually understandable. Microsoft Word is used for preparing the final submission in a clean academic style with separate sections, figures, and explanations. Together, these tools form a simple but complete data science workflow.")

    page_break(doc)

    add_page_heading(doc, "Methodology")
    methodology_paragraphs = [
        "Step 1: Two separate arrays are created, each containing 180 telecom usage values. This is done to represent data collected from different periods before combining them into one full dataset.",
        "Step 2: The two arrays are joined using NumPy concatenate(). This gives one continuous 360-value telecom usage array.",
        "Step 3: The one-dimensional array is reshaped into a 12 x 30 array using reshape(12, 30). In this format, each row shows one month and each column shows one day.",
        "Step 4: The reshaped array is converted into a Pandas DataFrame. Month names are added as row labels, and daily columns are added for better readability.",
        "Step 5: A new column named Monthly_Total is calculated by summing all daily values for each month.",
        "Step 6: The monthly totals are plotted using a line graph. This helps in understanding the monthly increase or decrease in telecom usage.",
    ]
    for text in methodology_paragraphs:
        add_paragraph(doc, text)
    add_paragraph(doc, "The methodology is arranged in a logical order so that each operation prepares the dataset for the next stage. Concatenation is performed before reshaping because the full yearly data must exist as one sequence before it can be divided into monthly rows. Similarly, the DataFrame is created after reshaping because Pandas works best when the structure already reflects the intended row and column arrangement.")
    add_paragraph(doc, "This sequence is important in real data science work. Analysts often receive incomplete or unstructured data first, then perform cleaning and formatting before carrying out any meaningful summaries or visualizations. The method followed in this project reflects that practical workflow in a simple student-friendly example.")
    add_image(doc, diagram_path, width=6.2, caption="Figure 1: Diagrammatic representation of the telecom data workflow")

    page_break(doc)

    add_page_heading(doc, "Implementation")
    add_paragraph(doc, "The following Python code completes the full process of dataset creation, shape fixing, DataFrame conversion, monthly total calculation, and graph generation.")
    add_paragraph(doc, "The code is written in a simple and readable form so that each operation can be understood clearly. Comments are used to explain the purpose of the major steps. The program first creates the telecom data, then unifies it into one dataset, corrects the data shape, converts it into a DataFrame, calculates monthly totals, and finally saves the graph for report use.")
    add_code_block(doc, code_text)

    page_break(doc)

    add_page_heading(doc, "Output and Visualization")
    add_paragraph(doc, f"When the code is executed, the original telecom array is shown with shape {telecom_usage_360.shape}. After applying reshape(), the corrected dataset appears with shape {telecom_usage_12x30.shape}. The DataFrame preview confirms that the values are arranged month-wise and day-wise, which makes analysis easier.")
    add_paragraph(doc, "The console output also displays monthly totals for all 12 months. These totals help compare usage levels across the year. The generated line graph clearly shows how telecom usage changes from month to month.")
    add_paragraph(doc, "The line graph is particularly helpful because it converts tabular totals into a visual pattern. Rather than reading one month at a time, the viewer can immediately notice where the usage begins to rise, where it stays stable, and where it reaches its highest level. In academic reporting, this visual support strengthens the explanation because the reader can connect the numbers with the trend more easily.")
    add_paragraph(doc, "The DataFrame preview is also an important output because it proves that the shape correction has been successful. Instead of an unclear one-dimensional list, the values are displayed in a structured table with month names and daily columns. This formatted output is exactly the type of representation that helps during further analysis or decision-making.")
    add_image(doc, output_path, width=5.8, caption="Figure 2: Program output summary")
    add_image(doc, preview_path, width=6.4, caption="Figure 3: DataFrame preview")
    add_image(doc, chart_path, width=6.3, caption="Figure 4: Monthly telecom usage line graph")
    add_paragraph(doc, "Suggested screenshots for submission: the full code window, the console output showing array shapes, the DataFrame preview, and the final line graph. These screenshots can be kept as supporting evidence if your faculty asks for execution proof in addition to the embedded figures.")

    page_break(doc)

    add_page_heading(doc, "Analysis and Discussion")
    add_paragraph(doc, "The corrected data shape makes the telecom dataset much easier to understand. Instead of one long array, the data now appears as a meaningful table of months and days. This allows direct comparison between months and supports simple trend analysis. Once the data is displayed in this form, it becomes possible to identify which months are relatively light, which months are stable, and which months place greater pressure on the network.")
    add_paragraph(doc, f"The monthly totals show that usage is lower in the first half of the year and higher in the later months. In this sample output, {peak_month} records the highest total usage of {peak_total}, while {min_month} records the lowest total usage of {min_total}. The graph shows a clear jump in network load from July onward, which suggests heavier traffic in the second half of the year. This pattern is useful because it provides a direct link between numerical totals and practical telecom demand.")
    add_paragraph(doc, "In simple terms, peak usage means more pressure on the telecom network. Higher usage months may require more bandwidth, stronger server support, or improved traffic handling methods. Lower usage months, on the other hand, may indicate periods where maintenance or optimization tasks can be planned more safely.")
    add_paragraph(doc, "The analysis also highlights the value of combining NumPy and Pandas in one workflow. NumPy efficiently handles shape transformation, while Pandas makes the result easier to read and summarize. Without these steps, identifying usage trends from a raw 360-length array would be slow and less intuitive. Therefore, the discussion confirms that proper data structure is a necessary foundation for network usage interpretation.")

    page_break(doc)

    add_page_heading(doc, "Conclusion")
    add_paragraph(doc, "This project shows how NumPy and Pandas can be used to fix the shape of telecom data in a simple and practical way. A 360-value dataset was created, merged with concatenate(), reshaped into a 12 x 30 structure, converted into a DataFrame, and analyzed through monthly totals and a line graph. Each stage of the project contributes to turning raw numerical values into meaningful information.")
    add_paragraph(doc, "The assignment also demonstrates that data formatting is not a minor step but an important part of analysis. Before any graph or summary can be trusted, the data must first be arranged correctly. Once the shape is fixed, the same values become easier to read, compare, and interpret. This is a useful lesson for data science students because many real problems begin with unstructured or poorly arranged data.")
    add_paragraph(doc, "In real telecom systems, properly structured data helps engineers and analysts monitor demand, identify peak traffic periods, and improve service quality. By understanding which months carry the highest network load, telecom providers can plan upgrades, reduce congestion, balance resources more effectively, and maintain a better experience for customers. This gives the mini-project a clear real-world relevance beyond classroom practice.")

    page_break(doc)

    add_page_heading(doc, "References")
    references = [
        "Wes McKinney, Python for Data Analysis, O'Reilly Media.",
        "Travis E. Oliphant, Guide to NumPy, Trelgol Publishing.",
        "Jake VanderPlas, Python Data Science Handbook, O'Reilly Media.",
        "Matplotlib Development Team, Matplotlib Documentation, official project documentation.",
    ]
    for ref in references:
        add_numbered(doc, ref)

    output_docx = ROOT / "telecom_data_shape_fix_report_expanded.docx"
    doc.save(output_docx)
    return output_docx


def main():
    telecom_usage_360, telecom_usage_12x30, telecom_df = generate_dataset()
    chart_path = save_line_chart(telecom_df)
    preview_path = save_dataframe_preview(telecom_df)
    output_path = save_console_output(telecom_usage_360, telecom_usage_12x30, telecom_df)
    diagram_path = save_workflow_diagram()
    docx_path = build_document(
        telecom_usage_360,
        telecom_usage_12x30,
        telecom_df,
        chart_path,
        preview_path,
        output_path,
        diagram_path,
    )
    print(f"Report generated: {docx_path}")


if __name__ == "__main__":
    main()
